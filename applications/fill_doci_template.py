"""Fill the official Elsevier Declaration of Interests template.

Marks the first option ("no competing interests") and adds the author
block at the end. Re-runs idempotently on the freshly-downloaded file.
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "applications/jmst_submission/DeclarationOfInterests.docx"
FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)


def style(run, size=11, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = BLACK


def main():
    doc = Document(str(TARGET))

    # Mark the first option (paragraph 6 in the template) by prepending [X].
    target_para = doc.paragraphs[6]
    if not target_para.text.strip().startswith("[X]"):
        # Prepend "[X] " by inserting a new run at the front
        # python-docx doesn't easily insert a new run at a specific position,
        # so we update the existing first run's text.
        if target_para.runs:
            first_run = target_para.runs[0]
            first_run.text = "[X] " + first_run.text.lstrip()
            # Ensure the leading marker is bold black
            first_run.bold = True
            first_run.font.color.rgb = BLACK
        else:
            run = target_para.add_run("[X] ")
            style(run, bold=True)

    # Walk every existing run and force black + named font to be safe
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.color.rgb = BLACK
            r.font.name = FONT

    # Append author block at the end
    def add(text, bold=False, italic=False, size=11):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        style(run, size=size, bold=bold, italic=italic)
        return p

    add("")
    add("Author name", bold=True)
    add("Sunwoo Lee")
    add("")
    add("Affiliation", bold=True)
    add("Independent researcher, South Korea")
    add("")
    add("ORCID", bold=True)
    add("0009-0004-9159-367X")
    add("")
    add("Date", bold=True)
    add("April 2026")
    add("")
    add("Signature", bold=True)
    add("Sunwoo Lee", italic=True)

    doc.save(str(TARGET))
    print(f"populated {TARGET}")


if __name__ == "__main__":
    main()
