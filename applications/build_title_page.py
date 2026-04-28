"""Render the JMST manuscript title page as a single Word document.

Includes: title, complete author list with affiliations, corresponding-author
contact block, abstract, and keywords. This is the file Editorial Manager
asks for under "Manuscript Title Page" — its data populates the metadata
fields later in the submission process.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches


HERE = Path(__file__).resolve().parents[1]
OUT = HERE / "applications/jmst_submission/TitlePage.docx"
FONT = "Times New Roman"


def style_run(run, size=12, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_centered(doc, text, size=12, bold=False, italic=False, space_after=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    style_run(p.add_run(text), size=size, bold=bold, italic=italic)
    return p


def add_left(doc, text, size=12, bold=False, italic=False,
             justified=True, space_after=4, line_spacing=1.15,
             space_before=0):
    p = doc.add_paragraph()
    p.alignment = (WD_ALIGN_PARAGRAPH.JUSTIFY if justified
                   else WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    style_run(p.add_run(text), size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, size=12, bold=True, space_before=10, space_after=2,
                justified=False):
    return add_left(doc, text, size=size, bold=bold, justified=justified,
                    space_before=space_before, space_after=space_after)


def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title
    add_centered(
        doc,
        "Cubic-Equivariant k-Space Convolution and a Differentiable "
        "Bond-Valence Field for Ionic-Conductivity Prediction in "
        "Solid-State Electrolytes",
        size=14, bold=True, space_after=14,
    )

    # Author block (single author, with superscript on name + affiliation marker)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Sunwoo Lee")
    style_run(r, size=12, bold=True)
    s = p.add_run("¹*")
    style_run(s, size=10)

    add_centered(doc, "(corresponding author)", size=10, italic=True, space_after=12)

    # Affiliations block
    add_heading(doc, "Affiliations", size=11, space_before=0)
    add_left(
        doc,
        "¹ Independent researcher, South Korea.",
        size=11, justified=False, line_spacing=1.15, space_after=4,
    )

    # Corresponding author contact details
    add_heading(doc, "Corresponding author", size=11)
    add_left(
        doc,
        "* Sunwoo Lee\n"
        "Independent researcher, South Korea\n"
        "ORCID: 0009-0004-9159-367X\n"
        "Email: lee.11539@buckeyemail.osu.edu",
        size=11, justified=False, line_spacing=1.2, space_after=4,
    )

    # Abstract
    add_heading(doc, "Abstract", size=12)
    abstract = (
        "Predicting the room-temperature ionic conductivity σ of "
        "solid-state lithium-ion electrolytes from crystal structure "
        "remains a bottleneck for materials discovery. I introduce two "
        "architectural primitives. k-SEC is a neural encoder whose "
        "feature maps live in reciprocal space throughout the network "
        "and whose cubic-group equivariance is enforced by construction "
        "via cubic-harmonic directional filters and a cross-shell gated "
        "attention modulated only by k-magnitudes. DBVF (Differentiable "
        "Bond-Valence Field) embeds Brown’s bond-valence sum into the "
        "network as a learnable module — the first treatment of Brown’s "
        "tabulated parameters as trainable parameters of a neural "
        "network. On the OBELiX benchmark (281 labelled solid "
        "electrolytes, 5-fold cross-validation × 5 seeds) the joint "
        "model attains MAE 1.047 [95 % CI 0.925–1.183] standalone and "
        "0.980 in a stacked ensemble with a gradient-boosted-tree "
        "baseline; Spearman ρ = 0.78 with 70 % top-10 precision. An "
        "unsupervised screen of 18,574 Li-containing Materials Project "
        "crystals identifies all four known fast-Li-ion conductor "
        "families in its top-15. A control experiment shows the DBVF "
        "gain is realised only through end-to-end training; a "
        "learning-curve experiment shows a 45 % gap reduction to the "
        "tabular baseline as training samples grow from 90 to 135 per "
        "fold, before plateauing at OBELiX scale."
    )
    add_left(doc, abstract, size=11, justified=True, line_spacing=1.2, space_after=8)

    # Keywords
    add_heading(doc, "Keywords", size=12)
    add_left(
        doc,
        "Solid electrolytes; Lithium-ion conductivity; Machine "
        "learning; Bond-valence model; Equivariant neural networks; "
        "OBELiX benchmark.",
        size=11, justified=False, line_spacing=1.15, space_after=4,
    )

    doc.save(str(OUT))
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
