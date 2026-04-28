"""Force every run of text in a .docx to RGB(0,0,0) black.

Pandoc's docx output uses styled hyperlinks (blue), accented heading
styles, and sometimes coloured table headers. Submission guidelines
require black-only text, so we walk every paragraph (including tables
and nested elements) and set the run's font color explicitly.

Usage:
    python applications/force_black_text.py applications/jmst_submission/*.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor


BLACK = RGBColor(0, 0, 0)


def force_run_black(run):
    """Set font color and clear any style-defined highlight on a run."""
    run.font.color.rgb = BLACK
    # Strip any direct fill/highlight on the run XML
    rpr = run._element.find(qn("w:rPr"))
    if rpr is not None:
        for tag in ("w:highlight", "w:shd"):
            for el in rpr.findall(qn(tag)):
                rpr.remove(el)


def walk_paragraphs(paragraphs):
    for p in paragraphs:
        for run in p.runs:
            force_run_black(run)


def walk_tables(tables):
    for table in tables:
        # Optional: force table border color to black
        tblPr = table._element.find(qn("w:tblPr"))
        for cell in (c for row in table.rows for c in row.cells):
            walk_paragraphs(cell.paragraphs)
            walk_tables(cell.tables)
        # Override default table styling colors
        for tc in table._element.iter(qn("w:tc")):
            shd = tc.find(qn("w:tcPr") + "/" + qn("w:shd"))
            if shd is not None:
                shd.set(qn("w:fill"), "FFFFFF")


def force_styles_black(doc):
    """Override the document's named styles so anything using them is black."""
    styles_xml = doc.styles.element
    for clr in styles_xml.iter(qn("w:color")):
        clr.set(qn("w:val"), "000000")
    # Hyperlink style in particular
    for st in styles_xml.iter(qn("w:style")):
        sid = st.get(qn("w:styleId"))
        if sid and "Hyperlink" in sid:
            rpr = st.find(qn("w:rPr"))
            if rpr is None:
                continue
            for clr in rpr.findall(qn("w:color")):
                clr.set(qn("w:val"), "000000")


def process(path: Path):
    doc = Document(str(path))
    walk_paragraphs(doc.paragraphs)
    walk_tables(doc.tables)
    # Headers / footers, if any
    for section in doc.sections:
        for header in (section.header, section.first_page_header,
                       section.even_page_header):
            walk_paragraphs(header.paragraphs)
            walk_tables(header.tables)
        for footer in (section.footer, section.first_page_footer,
                       section.even_page_footer):
            walk_paragraphs(footer.paragraphs)
            walk_tables(footer.tables)
    force_styles_black(doc)
    doc.save(str(path))
    print(f"  blackened: {path}")


def main():
    args = sys.argv[1:]
    if not args:
        print("usage: python force_black_text.py <files-or-dir>")
        return
    paths = []
    for a in args:
        p = Path(a)
        if p.is_dir():
            paths.extend(p.glob("*.docx"))
        elif p.suffix.lower() == ".docx":
            paths.append(p)
    if not paths:
        print("no .docx files matched")
        return
    for p in paths:
        process(p)


if __name__ == "__main__":
    main()
