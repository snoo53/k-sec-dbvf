"""Build the Declaration of Competing Interests document for JMST.

Elsevier's standard declaration. Since there are no conflicts to declare,
this is a short, signed statement.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches


HERE = Path(__file__).resolve().parents[1]
OUT = HERE / "applications/jmst_submission/DeclarationOfInterests.docx"
FONT = "Times New Roman"


def style(run, size=12, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_centered(doc, text, size=12, bold=False, italic=False, space_after=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    style(p.add_run(text), size=size, bold=bold, italic=italic)


def add_left(doc, text, size=12, bold=False, italic=False, space_after=4,
             line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    style(p.add_run(text), size=size, bold=bold, italic=italic)


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1.0); s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)

    add_centered(
        doc,
        "Declaration of Competing Interests",
        size=14, bold=True, space_after=14,
    )

    add_left(
        doc,
        "Manuscript title: Cubic-Equivariant k-Space Convolution and a "
        "Differentiable Bond-Valence Field for Ionic-Conductivity "
        "Prediction in Solid-State Electrolytes",
        size=11, italic=True, space_after=14,
    )

    add_left(
        doc,
        "The author whose name is listed immediately below certifies "
        "that he has no affiliations with or involvement in any "
        "organization or entity with any financial interest (such as "
        "honoraria; educational grants; participation in speakers' "
        "bureaus; membership, employment, consultancies, stock ownership, "
        "or other equity interest; and expert testimony or patent-"
        "licensing arrangements), or non-financial interest (such as "
        "personal or professional relationships, affiliations, knowledge "
        "or beliefs) in the subject matter or materials discussed in "
        "this manuscript.",
        size=11, space_after=18,
    )

    add_left(doc, "Author name", size=11, bold=True, space_after=2)
    add_left(doc, "Sunwoo Lee", size=11, space_after=14)

    add_left(doc, "Affiliation", size=11, bold=True, space_after=2)
    add_left(doc, "Independent researcher, South Korea", size=11, space_after=14)

    add_left(doc, "ORCID", size=11, bold=True, space_after=2)
    add_left(doc, "0009-0004-9159-367X", size=11, space_after=14)

    add_left(doc, "Date", size=11, bold=True, space_after=2)
    add_left(doc, "April 2026", size=11, space_after=14)

    add_left(doc, "Signature", size=11, bold=True, space_after=2)
    add_left(doc, "Sunwoo Lee", size=11, italic=True, space_after=4)

    doc.save(str(OUT))
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
